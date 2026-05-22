import io
import re
from pathlib import Path

try:
    import fitz
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None


def parse_uploaded_file(content: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(content)
    elif suffix == ".docx":
        return _parse_docx(content)
    elif suffix == ".txt":
        return _parse_txt(content)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _parse_pdf(content: bytes) -> str:
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) is not installed")
    doc = fitz.open(stream=content, filetype="pdf")
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    text = "\n".join(text_parts)
    return _clean_text(text)


def _parse_docx(content: bytes) -> str:
    if Document is None:
        raise ImportError("python-docx is not installed")
    doc = Document(io.BytesIO(content))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)
    text = "\n".join(text_parts)
    return _clean_text(text)


def _parse_txt(content: bytes) -> str:
    raw = content.decode("utf-8", errors="replace")
    return _clean_text(raw)


def _clean_text(text: str) -> str:
    text = re.sub(r"\u0000", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()
